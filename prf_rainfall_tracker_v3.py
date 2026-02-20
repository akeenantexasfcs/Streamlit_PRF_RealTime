"""
PRF Live Rainfall Tracker
Streamlit in Snowflake Application
Texas Farm Credit

v3.0 — Client Policies + Indemnity Estimates + Grid-County Separation
"""

import streamlit as st
import plotly.graph_objects as go
import pandas as pd
import numpy as np
import io
import datetime
from decimal import Decimal, ROUND_HALF_UP
from snowflake.snowpark.context import get_active_session

session = get_active_session()

# ─── Branding ───
FC_GREEN = "#5E9732"
FC_SLATE = "#5B707F"
FC_RUST  = "#9D5F58"
FC_CREAM = "#F5F1E8"
FC_AMBER = "#C4952B"
FC_DARK  = "#2d3a2e"

# ─── Interval Configuration ───
INTERVAL_CONFIG = {
    "625 — Jan-Feb": {
        "code": "625", "name": "Jan-Feb",
        "start": "01-01", "end": "02-28",
        "total_days": 59, "label": "January 1 – February 28"
    },
    "626 — Feb-Mar": {
        "code": "626", "name": "Feb-Mar",
        "start": "02-01", "end": "03-31",
        "total_days": 59, "label": "February 1 – March 31"
    },
}

# ─── Interval Name ↔ Code Mapping (all 11 PRF intervals) ───
INTERVAL_NAME_TO_CODE = {
    "Jan-Feb": "625", "Feb-Mar": "626", "Mar-Apr": "627",
    "Apr-May": "628", "May-Jun": "629", "Jun-Jul": "630",
    "Jul-Aug": "631", "Aug-Sep": "632", "Sep-Oct": "633",
    "Oct-Nov": "634", "Nov-Dec": "635",
}
INTERVAL_CODE_TO_NAME = {v: k for k, v in INTERVAL_NAME_TO_CODE.items()}


# ═══════════════════════════════════════════
# STYLES
# ═══════════════════════════════════════════

st.markdown(f"""
<style>
    .main .block-container,
    .appview-container .main .block-container,
    section[data-testid="stMainBlockContainer"] {{
        max-width: 100% !important;
        width: 100% !important;
        padding: 1rem 2rem !important;
    }}
    div[data-testid="stMetric"] {{
        background: {FC_CREAM};
        border: 2px solid {FC_GREEN};
        padding: 20px;
        border-radius: 10px;
        text-align: center;
    }}
    div[data-testid="stMetric"] label {{ 
        color: {FC_SLATE} !important; 
        font-size: 1rem !important;
        font-weight: 700 !important;
        letter-spacing: 0.5px !important;
    }}
    div[data-testid="stMetric"] [data-testid="stMetricValue"] {{ 
        color: #2d3a2e !important; 
        font-size: 2rem !important;
        font-weight: 800 !important;
    }}
    .signal-indemnity {{ 
        background: {FC_GREEN}; color: white; padding: 8px 20px; 
        border-radius: 8px; font-weight: 700; font-size: 1rem;
        text-align: center; margin-top: 4px;
    }}
    .signal-ok {{ 
        background: {FC_SLATE}; color: white; padding: 8px 20px; 
        border-radius: 8px; font-weight: 700; font-size: 1rem;
        text-align: center; margin-top: 4px;
    }}
    .indemnity-card {{
        border-radius: 10px; padding: 20px; text-align: center;
        font-weight: 700; color: white;
    }}
    section[data-testid="stSidebar"] {{ min-width: 320px; max-width: 380px; }}
    div[data-testid="stDataFrame"],
    div[data-testid="stDataFrame"] > div {{
        width: 100% !important;
    }}
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════
# CALCULATION FUNCTIONS (Decimal arithmetic — matches CvC app)
# ═══════════════════════════════════════════

def round_half_up(value, decimals=2):
    """Round using 'round half up' to match PRF official tool."""
    if value is None or (isinstance(value, float) and (np.isnan(value) or np.isinf(value))):
        return 0.0
    d = Decimal(str(value))
    if decimals == 0:
        quantize_to = Decimal('1')
    else:
        quantize_to = Decimal('0.' + '0' * (decimals - 1) + '1')
    return float(d.quantize(quantize_to, rounding=ROUND_HALF_UP))


def calculate_shortfall_decimal(trigger_level, index_value):
    """Calculate shortfall using Decimal to avoid IEEE 754 errors."""
    trigger_dec = Decimal(str(trigger_level))
    index_dec = Decimal(str(index_value))
    if index_dec >= trigger_dec:
        return Decimal('0')
    return (trigger_dec - index_dec) / trigger_dec


def calculate_protection_decimal(county_base_value, coverage_level, productivity_factor, insurable_interest=1.0):
    """Calculate dollar protection per acre with Decimal precision."""
    cbv  = Decimal(str(county_base_value))
    cov  = Decimal(str(coverage_level))
    prod = Decimal(str(productivity_factor))
    ins  = Decimal(str(insurable_interest))
    result = cbv * cov * prod * ins
    return float(result.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP))


def calculate_position_indemnity(final_index, coverage_level, county_base_value,
                                  productivity_factor, insurable_interest,
                                  acres, allocation_pct):
    """
    Calculate indemnity for one position for one interval.
    Matches CvC backtest Decimal arithmetic exactly.
    
    Args:
        final_index: Projected rainfall index value (0-150+)
        coverage_level: As decimal (e.g. 0.75)
        county_base_value: Dollar CBV
        productivity_factor: e.g. 1.35
        insurable_interest: e.g. 1.0
        acres: Total insured acres for this position
        allocation_pct: Percent allocated to this interval (e.g. 28 for 28%)
    
    Returns:
        Integer indemnity payment in dollars
    """
    trigger = coverage_level * 100  # 0.75 → 75.0

    if final_index >= trigger:
        return 0

    # Protection per acre (CBV × coverage × productivity × insurable interest)
    protection_per_acre = calculate_protection_decimal(
        county_base_value, coverage_level, productivity_factor, insurable_interest
    )
    total_protection = protection_per_acre * acres

    # Interval protection = Round(total_protection × allocation%)
    alloc_dec = Decimal(str(allocation_pct)) / Decimal('100')
    total_prot_dec = Decimal(str(total_protection))
    interval_protection_dec = (total_prot_dec * alloc_dec).quantize(
        Decimal('1'), rounding=ROUND_HALF_UP
    )
    interval_protection = int(interval_protection_dec)

    # Shortfall and indemnity
    shortfall_dec = calculate_shortfall_decimal(trigger, final_index)
    raw_indemnity = shortfall_dec * Decimal(str(interval_protection))

    if raw_indemnity >= Decimal('0.01'):
        return int(raw_indemnity.quantize(Decimal('1'), rounding=ROUND_HALF_UP))
    return 0


# ═══════════════════════════════════════════
# CACHED DATA LOADERS
# ═══════════════════════════════════════════

@st.cache_data(ttl=3600)
def load_texas_grids(interval_code):
    return session.sql(f"""
        WITH grid_counties AS (
            SELECT TRY_TO_NUMBER(m.SUB_COUNTY_CODE) AS GRID_ID,
                   LISTAGG(DISTINCT c.COUNTY_NAME, ' / ') WITHIN GROUP (ORDER BY c.COUNTY_NAME) AS COUNTY_NAME
            FROM MAP_YTD m
            LEFT JOIN COUNTY_YTD c 
                ON c.STATE_CODE = '48' AND c.COUNTY_CODE = m.COUNTY_CODE
                AND c.REINSURANCE_YEAR = 2025 AND c.DELETED_DATE IS NULL
            WHERE m.INSURANCE_PLAN_CODE = '13' AND m.STATE_CODE = '48' AND m.DELETED_DATE IS NULL
            GROUP BY 1
        )
        SELECT n.GRID_ID, n.NORMAL_IN, n.CV_PCT, n.CONFIDENCE_TIER,
               gc.COUNTY_NAME, g.CENTER_LAT, g.CENTER_LON
        FROM PRF_GRID_NORMALS n
        JOIN grid_counties gc ON n.GRID_ID = gc.GRID_ID
        LEFT JOIN PRF_GRID_REFERENCE g ON g.GRIDCODE = n.GRID_ID
        WHERE n.INTERVAL_CODE = '{interval_code}'
        ORDER BY n.GRID_ID
    """).to_pandas()


@st.cache_data(ttl=600)
def load_rainfall(date_start, date_end):
    """Uses PRF_GRID_CPC_MAP for correct coordinate snapping."""
    return session.sql(f"""
        SELECT 
            m.GRID_ID,
            ROUND(SUM(r.PRECIP_IN), 4) AS RAIN_SO_FAR,
            COUNT(DISTINCT r.OBSERVATION_DATE) AS DAYS_COLLECTED,
            MAX(r.OBSERVATION_DATE) AS LAST_DAY,
            MIN(r.FILE_TYPE) AS FILE_TYPE
        FROM PRF_RAINFALL_REALTIME r
        JOIN PRF_GRID_CPC_MAP m
            ON ROUND(r.LATITUDE, 3) = m.CPC_LAT
            AND ROUND(r.LONGITUDE, 3) = m.CPC_LON
        WHERE r.OBSERVATION_DATE BETWEEN '{date_start}' AND '{date_end}'
        GROUP BY 1
    """).to_pandas()


@st.cache_data(ttl=3600)
def load_client_list():
    """Load available client policies."""
    try:
        df = session.sql("""
            SELECT CLIENT_NAME, CROP_YEAR, POLICY_VERSION,
                   ARRAY_SIZE(POLICY_JSON:positions) AS POSITION_COUNT,
                   POLICY_JSON:defaults:coverage_level::FLOAT AS COVERAGE_LEVEL
            FROM PRF_CLIENT_POLICIES
            ORDER BY CLIENT_NAME, CROP_YEAR DESC
        """).to_pandas()
        return df
    except Exception:
        return pd.DataFrame()


@st.cache_data(ttl=600)
def load_client_policy(client_name, crop_year):
    """Load and flatten a client's policy JSON into a positions DataFrame."""
    df = session.sql(f"""
        SELECT
            p.value:grid_id::INT              AS GRID_ID,
            p.value:county::VARCHAR           AS COUNTY,
            p.value:state::VARCHAR            AS STATE,
            p.value:acres::INT                AS ACRES,
            p.value:county_base_value::FLOAT  AS COUNTY_BASE_VALUE,
            p.value:allocation                AS ALLOCATION_JSON,
            pol.POLICY_JSON:defaults:coverage_level::FLOAT       AS COVERAGE_LEVEL,
            pol.POLICY_JSON:defaults:productivity_factor::FLOAT  AS PRODUCTIVITY_FACTOR,
            pol.POLICY_JSON:defaults:insurable_interest::FLOAT   AS INSURABLE_INTEREST
        FROM PRF_CLIENT_POLICIES pol,
             LATERAL FLATTEN(input => pol.POLICY_JSON:positions) p
        WHERE pol.CLIENT_NAME = '{client_name}'
          AND pol.CROP_YEAR = {crop_year}
    """).to_pandas()

    # Parse allocation JSON string into Python dicts
    import json
    def parse_alloc(val):
        if isinstance(val, str):
            return json.loads(val)
        elif isinstance(val, dict):
            return val
        return {}

    df['ALLOCATION'] = df['ALLOCATION_JSON'].apply(parse_alloc)
    return df


# ═══════════════════════════════════════════
# TRACKER & GAUGE FUNCTIONS
# ═══════════════════════════════════════════

def build_tracker(grids_df, rain_df, coverage_level, total_days):
    merged = grids_df.merge(rain_df, on="GRID_ID", how="inner")
    merged["PARTIAL_INDEX"] = (merged["RAIN_SO_FAR"] / merged["NORMAL_IN"] * 100).round(1)
    merged["DAILY_RATE"] = merged["RAIN_SO_FAR"] / merged["DAYS_COLLECTED"]
    merged["PROJECTED_RAIN"] = (merged["DAILY_RATE"] * total_days).round(4)
    merged["PROJECTED_INDEX"] = (merged["PROJECTED_RAIN"] / merged["NORMAL_IN"] * 100).round(1)
    trigger = coverage_level
    merged["SIGNAL"] = merged["PROJECTED_INDEX"].apply(
        lambda idx: "LIKELY INDEMNITY" if idx < trigger else "OK"
    )
    return merged


def create_gauge(grid_id, projected_index, partial_index, signal,
                 rain_so_far, normal_in, days, total_days,
                 coverage_level, county_name=None):

    bar_color = FC_GREEN if signal == "LIKELY INDEMNITY" else FC_SLATE
    trigger = coverage_level
    max_range = max(150, projected_index + 20)
    pct_through = round(days / total_days * 100)

    county_str = ""
    if county_name and pd.notna(county_name):
        county_str = f"  ·  {county_name}"

    fig = go.Figure()
    fig.add_trace(go.Indicator(
        mode="gauge+number",
        value=projected_index,
        number={
            "valueformat": ".1f",
            "font": {"size": 64, "color": "#2d3a2e", "family": "Arial Black"},
        },
        title={
            "text": (
                f"<b style='font-size:20px'>Grid {grid_id}</b>"
                f"<span style='font-size:13px;color:{FC_SLATE}'>{county_str}</span>"
                f"<br>"
                f"<span style='font-size:14px;color:{FC_SLATE}'>"
                f"Rain: <b>{rain_so_far:.2f}\"</b> of {normal_in:.1f}\" normal"
                f"  ·  {days}/{total_days} days ({pct_through}%)"
                f"  ·  Coverage: {coverage_level}%"
                f"</span>"
            ),
            "font": {"size": 14, "color": "#2d3a2e"},
        },
        gauge={
            "axis": {
                "range": [0, max_range], "tickwidth": 2,
                "tickcolor": FC_SLATE, "tickfont": {"color": FC_SLATE, "size": 14},
                "dtick": 25,
            },
            "bar": {"color": bar_color, "thickness": 0.75},
            "bgcolor": "#e8e4dd", "borderwidth": 0,
            "steps": [
                {"range": [0, trigger], "color": "rgba(94, 151, 50, 0.28)"},
                {"range": [trigger, max_range], "color": "rgba(91, 112, 127, 0.08)"},
            ],
            "threshold": {
                "line": {"color": "#2d3a2e", "width": 5},
                "thickness": 0.9, "value": partial_index,
            },
        },
        domain={"x": [0.15, 0.85], "y": [0, 0.85]},
    ))

    fig.update_layout(
        height=300,
        margin=dict(l=20, r=20, t=100, b=10),
        paper_bgcolor=FC_CREAM, plot_bgcolor=FC_CREAM,
        font={"color": "#2d3a2e"},
    )
    return fig


# ═══════════════════════════════════════════
# INDEMNITY SCENARIO ENGINE
# ═══════════════════════════════════════════

def compute_indemnity_scenarios(positions_df, rain_df, normals_df, interval_name, total_days):
    """
    For each client position with allocation for the given interval,
    compute three indemnity scenarios: pessimistic, current trend, optimistic.
    
    Returns a DataFrame with one row per position.
    """
    rows = []

    for _, pos in positions_df.iterrows():
        alloc = pos['ALLOCATION']
        alloc_pct = alloc.get(interval_name, 0)
        if alloc_pct == 0:
            continue  # Position has no allocation for this interval

        grid_id = pos['GRID_ID']

        # Look up rainfall for this grid
        rain_row = rain_df[rain_df['GRID_ID'] == grid_id]
        if rain_row.empty:
            rain_so_far = 0.0
            days_collected = 0
        else:
            rain_so_far = float(rain_row.iloc[0]['RAIN_SO_FAR'])
            days_collected = int(rain_row.iloc[0]['DAYS_COLLECTED'])

        # Look up normal for this grid
        normal_row = normals_df[normals_df['GRID_ID'] == grid_id]
        if normal_row.empty:
            continue  # Can't compute without normal
        normal_in = float(normal_row.iloc[0]['NORMAL_IN'])
        if normal_in <= 0:
            continue

        # Current partial index
        partial_index = round((rain_so_far / normal_in) * 100, 1)

        # ─── Three Scenarios ───
        # From PRF policyholder perspective: more payout = better
        # Optimistic = highest payout (zero rain remaining)
        # Pessimistic = lowest payout (normal pace remaining)
        remaining_days = max(0, total_days - days_collected)

        # Optimistic (for PRF): zero rain remaining → max shortfall → max payout
        optimistic_rain = rain_so_far
        optimistic_index = round((optimistic_rain / normal_in) * 100, 1)

        # Current trend: linear extrapolation
        if days_collected > 0:
            daily_rate = rain_so_far / days_collected
            trend_rain = round(daily_rate * total_days, 4)
        else:
            trend_rain = 0.0
        trend_index = round((trend_rain / normal_in) * 100, 1)

        # Pessimistic (for PRF): remaining days at normal rate → min shortfall → min payout
        normal_daily_rate = normal_in / total_days
        pessimistic_rain = rain_so_far + (normal_daily_rate * remaining_days)
        pessimistic_index = round((pessimistic_rain / normal_in) * 100, 1)

        # ─── Dollar Indemnities ───
        cov   = pos['COVERAGE_LEVEL']
        cbv   = pos['COUNTY_BASE_VALUE']
        pf    = pos['PRODUCTIVITY_FACTOR']
        ii    = pos['INSURABLE_INTEREST']
        acres = pos['ACRES']

        protection_per_acre = calculate_protection_decimal(cbv, cov, pf, ii)
        total_protection = protection_per_acre * acres

        indem_pessimistic = calculate_position_indemnity(
            pessimistic_index, cov, cbv, pf, ii, acres, alloc_pct)
        indem_trend = calculate_position_indemnity(
            trend_index, cov, cbv, pf, ii, acres, alloc_pct)
        indem_optimistic = calculate_position_indemnity(
            optimistic_index, cov, cbv, pf, ii, acres, alloc_pct)

        rows.append({
            'GRID_ID': grid_id,
            'COUNTY': pos['COUNTY'],
            'ACRES': acres,
            'ALLOC_PCT': alloc_pct,
            'CBV': cbv,
            'NORMAL_IN': normal_in,
            'RAIN_SO_FAR': rain_so_far,
            'DAYS': days_collected,
            'PARTIAL_INDEX': partial_index,
            'PESSIMISTIC_IDX': pessimistic_index,
            'TREND_IDX': trend_index,
            'OPTIMISTIC_IDX': optimistic_index,
            'PROTECTION': round(total_protection * alloc_pct / 100, 0),
            'INDEM_PESSIMISTIC': indem_pessimistic,
            'INDEM_TREND': indem_trend,
            'INDEM_OPTIMISTIC': indem_optimistic,
        })

    return pd.DataFrame(rows) if rows else pd.DataFrame()


# ═══════════════════════════════════════════
# WORD REPORT GENERATOR
# ═══════════════════════════════════════════

def _angle_for_value(val, max_range):
    """Convert an index value to a gauge arc angle (0=right, 180=left)."""
    return 180 - (val / max_range) * 180


def generate_gauge_png(grid_id, projected_index, partial_index, signal,
                       rain_so_far, normal_in, days, total_days,
                       coverage_level, county_name=""):
    """Render a single gauge chart as a PNG BytesIO buffer with bottom status banner."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import Wedge, FancyBboxPatch

    max_range = max(150, projected_index + 20)
    pct_through = round(days / total_days * 100) if total_days > 0 else 0

    fig, ax = plt.subplots(1, 1, figsize=(3.8, 3.1), facecolor=FC_CREAM)
    ax.set_xlim(-1.3, 1.3)
    ax.set_ylim(-0.52, 1.45)
    ax.set_aspect('equal')
    ax.axis('off')

    # Background arc
    ax.add_patch(Wedge((0, 0), 1.0, 0, 180, width=0.28, facecolor='#e8e4dd', edgecolor='none'))
    # Indemnity zone (0 -> trigger)
    trigger_angle = _angle_for_value(coverage_level, max_range)
    ax.add_patch(Wedge((0, 0), 1.0, trigger_angle, 180, width=0.28, facecolor='#5E973248', edgecolor='none'))
    # Value bar
    proj_angle = _angle_for_value(min(projected_index, max_range), max_range)
    bar_color = FC_GREEN if signal == "LIKELY INDEMNITY" else FC_SLATE
    ax.add_patch(Wedge((0, 0), 0.97, proj_angle, 180, width=0.22, facecolor=bar_color, edgecolor='none'))
    # Threshold needle (partial/current index)
    needle_angle = _angle_for_value(min(partial_index, max_range), max_range)
    rad = np.radians(needle_angle)
    ax.plot([0.7 * np.cos(rad), 1.02 * np.cos(rad)],
            [0.7 * np.sin(rad), 1.02 * np.sin(rad)],
            color=FC_DARK, linewidth=2.5, zorder=5)
    # Tick marks
    for tick_val in range(0, int(max_range) + 1, 25):
        a = np.radians(_angle_for_value(tick_val, max_range))
        ax.plot([1.02 * np.cos(a), 1.08 * np.cos(a)],
                [1.02 * np.sin(a), 1.08 * np.sin(a)], color=FC_SLATE, linewidth=0.8)
        ax.text(1.15 * np.cos(a), 1.15 * np.sin(a), str(tick_val),
                ha='center', va='center', fontsize=6, color=FC_SLATE, fontweight='bold')
    # Big number
    ax.text(0, 0.38, f"{projected_index:.1f}", ha='center', va='center',
            fontsize=28, fontweight='black', color=FC_DARK, fontfamily='sans-serif')
    # Title line 1
    county_str = f"  \u00b7  {county_name}" if county_name else ""
    ax.text(0, 1.42, f"Grid {grid_id}{county_str}",
            ha='center', va='center', fontsize=10, fontweight='bold', color=FC_DARK)
    # Title line 2
    ax.text(0, 1.28,
            f'Rain: {rain_so_far:.2f}" of {normal_in:.1f}" normal  \u00b7  '
            f'{days}/{total_days} days ({pct_through}%)  \u00b7  Coverage: {coverage_level}%',
            ha='center', va='center', fontsize=5.5, color=FC_SLATE)

    # Bottom Banner (green="LIKELY INDEMNITY", slate="OK")
    banner_color = FC_GREEN if signal == "LIKELY INDEMNITY" else FC_SLATE
    banner = FancyBboxPatch((-1.25, -0.52), 2.5, 0.28,
                            boxstyle="round,pad=0.03", facecolor=banner_color, edgecolor='none')
    ax.add_patch(banner)
    if signal == "LIKELY INDEMNITY":
        txt = (f"LIKELY INDEMNITY  \u2014  Current: {partial_index:.1f}  \u00b7  "
               f"Projected: {projected_index:.1f}  \u00b7  Trigger: {coverage_level}")
    else:
        txt = f"OK  \u2014  Current: {partial_index:.1f}  \u00b7  Projected: {projected_index:.1f}"
    ax.text(0, -0.38, txt, ha='center', va='center',
            fontsize=5.5, fontweight='bold', color='white')

    plt.tight_layout(pad=0.2)
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight', facecolor=FC_CREAM)
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_legend_png(coverage_level):
    """Render the gauge legend bar as a PNG BytesIO buffer."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    fig, ax = plt.subplots(1, 1, figsize=(7.0, 0.5), facecolor=FC_CREAM)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 1)
    ax.axis('off')

    ax.add_patch(FancyBboxPatch((0.2, 0.25), 0.6, 0.5, boxstyle="round,pad=0.05",
                                facecolor=FC_GREEN, edgecolor='none'))
    ax.text(1.0, 0.5, "Bar \u2014 Projected Final Index", va='center', fontsize=8,
            fontweight='bold', color=FC_DARK)

    ax.plot([3.9, 3.9], [0.15, 0.85], color=FC_DARK, linewidth=2.5)
    ax.text(4.1, 0.5, "Line \u2014 Current Estimated Index", va='center', fontsize=8,
            fontweight='bold', color=FC_DARK)

    ax.add_patch(FancyBboxPatch((7.2, 0.25), 0.6, 0.5, boxstyle="round,pad=0.05",
                                facecolor='#5E973248', edgecolor='#ccc', linewidth=0.5))
    ax.text(8.0, 0.5, f"Indemnity zone (below {coverage_level})", va='center', fontsize=8,
            color=FC_DARK)

    plt.tight_layout(pad=0.1)
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight', facecolor=FC_CREAM)
    plt.close(fig)
    buf.seek(0)
    return buf


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell (color_hex without #)."""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_green_divider(doc, space_before=0, space_after=0, sz=8):
    """Add a green bottom-border paragraph as visual divider.
    CRITICAL: pBdr must be inserted at position 0 in pPr (before spacing/jc)
    or OOXML validation will fail."""
    from docx.shared import Pt
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>'
                     f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="5E9732"/>'
                     f'</w:pBdr>')
    pPr.insert(0, pBdr)
    return p


def build_word_report(client_name, crop_year, interval_name, interval_code,
                      interval_label, coverage_level, total_days,
                      indem_df, display_df, all_positions_df=None,
                      estimated_premium=None):
    """
    Build a complete PRF Rainfall Tracker Word report.
    Returns BytesIO buffer containing the .docx file.

    Report structure:
    - Page 1: Cover (Texas Farm Credit branding, client, date, BETA badge)
    - Page 2: Executive Summary (metrics, allocation table, disclaimer)
    - Page 3: Interval Detail (scenario cards, position breakdown table)
    - Page 4+: Gauges (legend + 3-per-row grid, overflow to next page)
    - Final: Glossary (scenario definitions, PRF terms, beta notice)
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    _GREEN = RGBColor(0x5E, 0x97, 0x32)
    _SLATE = RGBColor(0x5B, 0x70, 0x7F)
    _RUST  = RGBColor(0x9D, 0x5F, 0x58)
    _AMBER = RGBColor(0xC4, 0x95, 0x2B)
    _DARK  = RGBColor(0x2D, 0x3A, 0x2E)
    _WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    _WARN  = RGBColor(0x85, 0x64, 0x04)

    doc = Document()

    # Page setup: US Letter, 0.75" margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.6)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.font.color.rgb = _DARK

    report_date = datetime.date.today().strftime("%B %d, %Y")
    days_in = int(display_df["DAYS_COLLECTED"].iloc[0]) if len(display_df) > 0 else 0

    # ═══════════════════════════════════════
    # PAGE 1: COVER
    # ═══════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Texas Farm Credit")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = _GREEN
    run.font.name = 'Arial'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRF Live Rainfall Tracker")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = _DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Reporting")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = _DARK

    add_green_divider(doc, space_before=12, space_after=12, sz=12)

    for label, value, val_color in [
        ("Client: ", client_name, _DARK),
        ("Report Generated: ", report_date, _DARK),
        ("Generated by: ", "Texas Farm Credit", _GREEN),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label)
        r1.font.size = Pt(13)
        r1.font.color.rgb = _SLATE
        r2 = p.add_run(value)
        r2.font.size = Pt(13)
        r2.font.bold = True
        r2.font.color.rgb = val_color

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BETA")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = _AMBER

    for txt in [
        "This report is in beta mode. All figures are estimates only.",
        "Estimates are based on real-time CPC rainfall data and linear projections.",
        "Final index values and indemnities are determined by RMA/USDA after interval close.",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = _SLATE

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("FOR INTERNAL USE ONLY")
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = _RUST

    # ═══════════════════════════════════════
    # PAGE 2: EXECUTIVE SUMMARY
    # ═══════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("Executive Summary")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = _DARK

    add_green_divider(doc)

    metrics = [
        ("Client", f"{client_name} ({crop_year})"),
        ("Coverage Level", f"{coverage_level}%"),
        ("Interval in Report", f"{interval_name} ({interval_code})"),
    ]
    if estimated_premium:
        metrics.insert(1, ("Estimated Premium Cost", f"${estimated_premium:,.0f}"))
    if all_positions_df is not None:
        total_acres = int(all_positions_df['ACRES'].sum())
        metrics.append(("Total Insured Acres", f"{total_acres:,}"))

    for label, value in metrics:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        r1.font.size = Pt(11)
        r1.font.color.rgb = _SLATE
        r2 = p.add_run(value)
        r2.font.size = Pt(11)
        r2.font.bold = True
        r2.font.color.rgb = _DARK

    # Allocation Table
    if all_positions_df is not None and not all_positions_df.empty:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Current Setup \u2014 Grid Allocation by Interval")
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = _DARK

        iv_names = ["Jan-Feb", "Feb-Mar", "Mar-Apr", "Apr-May", "May-Jun",
                     "Jun-Jul", "Jul-Aug", "Aug-Sep", "Sep-Oct", "Oct-Nov", "Nov-Dec"]
        cols = ["Grid"] + iv_names + ["Acres"]
        num_cols = len(cols)

        tbl = doc.add_table(rows=1, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True

        hdr = tbl.rows[0]
        for i, col_name in enumerate(cols):
            cell = hdr.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(col_name)
            run.font.size = Pt(7)
            run.font.bold = True
            run.font.color.rgb = _WHITE
            run.font.name = 'Arial'
            set_cell_shading(cell, "2D3A2E")

        for _, pos in all_positions_df.iterrows():
            alloc = pos.get('ALLOCATION', {})
            row_cells = tbl.add_row().cells
            row_cells[0].text = ""
            p = row_cells[0].paragraphs[0]
            run = p.add_run(f"{pos['GRID_ID']} ({pos.get('COUNTY', '')})")
            run.font.size = Pt(7)
            run.font.name = 'Arial'

            for j, iv_name in enumerate(iv_names):
                pct = alloc.get(iv_name, 0)
                cell = row_cells[j + 1]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"{pct:.0f}%" if pct > 0 else "-")
                run.font.size = Pt(7)
                run.font.name = 'Arial'

            cell = row_cells[num_cols - 1]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"{int(pos['ACRES']):,}")
            run.font.size = Pt(7)
            run.font.name = 'Arial'

    # Beta disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "\u26a0  IMPORTANT: All indemnity figures in this report are estimates based on "
        "current data and projections. Final values are determined by RMA/USDA after "
        "interval close. This report is in BETA and should not be used as the sole "
        "basis for financial decisions."
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = _WARN

    # ═══════════════════════════════════════
    # PAGE 3: INTERVAL DETAIL
    # ═══════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run(f"{interval_name} {crop_year}")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = _DARK

    p = doc.add_paragraph()
    run = p.add_run(f"Interval {interval_code}  \u00b7  {interval_label}")
    run.font.size = Pt(11)
    run.font.color.rgb = _SLATE

    add_green_divider(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Portfolio Indemnity Estimates")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = _DARK

    if not indem_df.empty:
        total_opt   = int(indem_df['INDEM_OPTIMISTIC'].sum())
        total_trend = int(indem_df['INDEM_TREND'].sum())
        total_pess  = int(indem_df['INDEM_PESSIMISTIC'].sum())

        # Scenario cards (3-column colored table)
        card_table = doc.add_table(rows=1, cols=3)
        card_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        scenarios = [
            ("OPTIMISTIC", "Zero rain remaining", total_opt, "5E9732"),
            ("CURRENT TREND", "Linear extrapolation", total_trend, "C4952B"),
            ("PESSIMISTIC", "Normal pace remaining", total_pess, "9D5F58"),
        ]

        for i, (label, subtitle, amount, bg_hex) in enumerate(scenarios):
            cell = card_table.rows[0].cells[i]
            set_cell_shading(cell, bg_hex)
            cell.text = ""
            for txt, sz, bld in [(label, 10, True), (subtitle, 8, False), (f"${amount:,}", 20, True)]:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(txt)
                run.font.size = Pt(sz)
                run.font.bold = bld
                run.font.color.rgb = _WHITE
                run.font.name = 'Arial'

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(
            f"* Estimates based on rainfall data through day {days_in} of {total_days}. "
            "See Glossary for scenario methodology."
        )
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = _SLATE

        # Position Breakdown Table (15 columns)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run("Position Breakdown")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = _DARK

        pos_cols = ["Grid", "County", "Acres", "Alloc%", "CBV", "Normal", "Rain",
                    "Days", "Opt Idx", "Trend Idx", "Pess Idx",
                    "Protection", "$ Opt", "$ Trend", "$ Pess"]

        pos_tbl = doc.add_table(rows=1, cols=len(pos_cols))
        pos_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        pos_tbl.autofit = True

        for i, col_name in enumerate(pos_cols):
            cell = pos_tbl.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(col_name)
            run.font.size = Pt(7)
            run.font.bold = True
            run.font.color.rgb = _WHITE
            run.font.name = 'Arial'
            set_cell_shading(cell, "2D3A2E")

        for _, r in indem_df.iterrows():
            vals = [
                str(int(r['GRID_ID'])), r['COUNTY'], f"{int(r['ACRES']):,}",
                f"{r['ALLOC_PCT']:.0f}%", f"${r['CBV']:.2f}",
                f"{r['NORMAL_IN']:.2f}\"", f"{r['RAIN_SO_FAR']:.2f}\"",
                str(int(r['DAYS'])),
                f"{r['OPTIMISTIC_IDX']:.1f}", f"{r['TREND_IDX']:.1f}", f"{r['PESSIMISTIC_IDX']:.1f}",
                f"${int(r['PROTECTION']):,}",
                f"${int(r['INDEM_OPTIMISTIC']):,}",
                f"${int(r['INDEM_TREND']):,}",
                f"${int(r['INDEM_PESSIMISTIC']):,}",
            ]
            row_cells = pos_tbl.add_row().cells
            for i, v in enumerate(vals):
                cell = row_cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(v)
                run.font.size = Pt(7)
                run.font.name = 'Arial'

        # Totals row
        total_acres = int(indem_df['ACRES'].sum())
        total_prot  = int(indem_df['PROTECTION'].sum())
        tot_vals = [
            "TOTAL", "", f"{total_acres:,}", "", "", "", "", "", "", "", "",
            f"${total_prot:,}", f"${total_opt:,}", f"${total_trend:,}", f"${total_pess:,}",
        ]
        row_cells = pos_tbl.add_row().cells
        for i, v in enumerate(tot_vals):
            cell = row_cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            run.font.size = Pt(7)
            run.font.bold = True
            run.font.name = 'Arial'
            set_cell_shading(cell, "F5F1E8")

    # ═══════════════════════════════════════
    # PAGE 4+: GAUGES
    # ═══════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("Projected Final Index")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = _DARK

    add_green_divider(doc)

    # Legend
    legend_buf = generate_legend_png(coverage_level)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(legend_buf, width=Inches(5.5))

    # Gauges: 3 per row, sorted by projected index ascending
    display_sorted = display_df.sort_values("PROJECTED_INDEX", ascending=True)
    gauge_rows_data = [row for _, row in display_sorted.iterrows()]

    COLS = 3
    for chunk_start in range(0, len(gauge_rows_data), COLS):
        chunk = gauge_rows_data[chunk_start:chunk_start + COLS]
        gauge_tbl = doc.add_table(rows=1, cols=COLS)
        gauge_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(chunk):
            buf = generate_gauge_png(
                grid_id=row["GRID_ID"],
                projected_index=row["PROJECTED_INDEX"],
                partial_index=row["PARTIAL_INDEX"],
                signal=row["SIGNAL"],
                rain_so_far=row["RAIN_SO_FAR"],
                normal_in=row["NORMAL_IN"],
                days=row["DAYS_COLLECTED"],
                total_days=total_days,
                coverage_level=coverage_level,
                county_name=row.get("COUNTY_NAME", ""),
            )
            cell = gauge_tbl.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(buf, width=Inches(2.2))

        for i in range(len(chunk), COLS):
            gauge_tbl.rows[0].cells[i].text = ""

    # ═══════════════════════════════════════
    # GLOSSARY
    # ═══════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("Glossary")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = _DARK

    add_green_divider(doc)

    glossary = [
        ("Optimistic Scenario (Green)",
         "Assumes zero additional rainfall for the remainder of the interval. "
         "This produces the lowest possible final index and therefore the highest "
         "potential indemnity payout."),
        ("Current Trend Scenario (Amber)",
         "Uses linear extrapolation of the observed daily rainfall rate to project "
         "the final index. Daily rate = total rainfall / collection days, projected "
         "over the full interval."),
        ("Pessimistic Scenario (Rust)",
         "Assumes remaining days receive rainfall at the historical normal daily rate. "
         "This produces the highest possible final index and therefore the lowest "
         "potential indemnity payout."),
        ("Rainfall Index",
         "The ratio of actual accumulated rainfall to normal rainfall for a grid and "
         "interval, expressed as a percentage. Index below coverage level triggers indemnity."),
        ("Coverage Level",
         "The index threshold below which an indemnity payment is triggered (e.g. 75% "
         "means index below 75 triggers a payment)."),
        ("County Base Value (CBV)",
         "Dollar value per acre established by RMA for each county/commodity/practice combination."),
        ("Protection",
         "Total insurance coverage: CBV x Coverage Level x Productivity Factor x "
         "Insurable Interest x Acres x Allocation %."),
        ("Indemnity",
         "Payment when final index < coverage level: (Coverage - Index) / Coverage x "
         "Protection. Calculated with Decimal arithmetic and Round Half Up to match "
         "official RMA methods."),
    ]

    for term, definition in glossary:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(term)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = _DARK

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(definition)
        run.font.size = Pt(9)
        run.font.color.rgb = _SLATE

    # Final beta notice
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "\u26a0  BETA NOTICE: This report and all calculations are in beta testing. "
        "Indemnity estimates are projections based on partial-interval rainfall data "
        "and should not be used as the sole basis for financial planning. Final index "
        "values and indemnity amounts are determined exclusively by the USDA Risk "
        "Management Agency (RMA) after each interval closes. Texas Farm Credit makes "
        "no guarantee of accuracy."
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = _WARN

    # Fix zoom (python-docx omits required percent attribute)
    settings = doc.settings.element
    zoom = settings.find(qn('w:zoom'))
    if zoom is not None and zoom.get(qn('w:percent')) is None:
        zoom.set(qn('w:percent'), '100')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════

with st.sidebar:
    st.markdown(f"### ⚙️ Controls")

    # ── INTERVAL SELECTOR ──
    interval_selection = st.selectbox(
        "📅 Interval",
        options=list(INTERVAL_CONFIG.keys()),
        index=0,
        help="Select the coverage interval to track"
    )
    interval = INTERVAL_CONFIG[interval_selection]

    # ── CLIENT POLICY SELECTOR ──
    st.divider()
    st.markdown("**📋 Client Policy**")

    client_list_df = load_client_list()
    client_mode = False
    client_positions_df = None
    policy_coverage = None

    if not client_list_df.empty:
        client_options = ["— Manual —"] + [
            f"{r['CLIENT_NAME']} ({int(r['CROP_YEAR'])})"
            for _, r in client_list_df.iterrows()
        ]
        client_selection = st.selectbox("Load Client", client_options, index=0)

        if client_selection != "— Manual —":
            client_mode = True
            # Parse selection back to name + year
            sel_row = client_list_df.iloc[client_options.index(client_selection) - 1]
            client_name = sel_row['CLIENT_NAME']
            crop_year = int(sel_row['CROP_YEAR'])
            policy_coverage = float(sel_row['COVERAGE_LEVEL'])

            client_positions_df = load_client_policy(client_name, crop_year)

            # Filter to positions with allocation for selected interval
            interval_name = interval['name']  # e.g. "Jan-Feb"
            mask = client_positions_df['ALLOCATION'].apply(
                lambda a: a.get(interval_name, 0) > 0
            )
            client_positions_df = client_positions_df[mask].copy()

            pos_ct = len(client_positions_df)
            total_ac = client_positions_df['ACRES'].sum() if pos_ct > 0 else 0
            st.caption(
                f"✅ **{client_name}** — {pos_ct} positions for {interval_name} "
                f"· {total_ac:,} acres"
            )
    else:
        st.caption("No client policies found. Using manual mode.")

    st.divider()

    # ── COVERAGE LEVEL ──
    cov_options = [90, 85, 80, 75, 70]
    if client_mode and policy_coverage:
        # Map policy decimal to display integer (0.75 → 75)
        policy_cov_int = int(policy_coverage * 100)
        default_idx = cov_options.index(policy_cov_int) if policy_cov_int in cov_options else 0
    else:
        default_idx = 0

    coverage_level = st.selectbox(
        "Coverage Level",
        options=cov_options,
        index=default_idx,
        help="Indemnity triggers below this index level"
    )

    st.divider()

    # ── GRID SELECTION (manual mode only) ──
    if not client_mode:
        grids_df = load_texas_grids(interval["code"])

        grids_df["LABEL"] = grids_df.apply(
            lambda r: f"{r['GRID_ID']} — {r['COUNTY_NAME']}"
            if pd.notna(r.get("COUNTY_NAME")) else str(r["GRID_ID"]), axis=1
        )
        label_to_id = dict(zip(grids_df["LABEL"], grids_df["GRID_ID"]))

        all_counties = set()
        for names in grids_df["COUNTY_NAME"].dropna():
            for c in names.split(" / "):
                all_counties.add(c.strip())

        selected_counties = st.multiselect("Filter by County", sorted(all_counties), default=[])

        if selected_counties:
            mask = grids_df["COUNTY_NAME"].apply(
                lambda x: any(c in str(x) for c in selected_counties) if pd.notna(x) else False
            )
            filtered_labels = grids_df[mask]["LABEL"].tolist()
        else:
            filtered_labels = grids_df["LABEL"].tolist()

        st.markdown("**Select Grids**")
        grid_entry = st.text_input(
            "Enter Grid IDs (comma separated)",
            placeholder="7929, 8230, 8231",
            help="Type grid IDs directly"
        )
        selected_labels = st.multiselect("Or pick from list", filtered_labels, default=[])

        st.divider()
        st.markdown("**Quick Select**")
        col1, col2 = st.columns(2)
        with col1:
            top_n = st.selectbox("Driest N", [10, 25, 50, "All"], index=0)
        with col2:
            show_all_likely = st.checkbox("Likely only", value=False)
    else:
        # In client mode, load grids for gauge display
        grids_df = load_texas_grids(interval["code"])
        grid_entry = None
        selected_labels = []
        top_n = 10
        show_all_likely = False

    st.divider()
    generate = st.button("🚀 Generate", type="primary", use_container_width=True)

    if client_mode:
        st.divider()
        st.markdown("**📄 Word Report**")
        generate_report = st.button("📄 Generate Report", use_container_width=True,
                                     help="Download Word doc with scenarios, tables, and gauges")
    else:
        generate_report = False

    st.divider()
    st.markdown(f"""
    <div style='font-size:11px;color:{FC_SLATE}'>
        <b>Sources:</b> CPC Gauge (RT) · 3yr normals · Linear projection<br>
        <b>Interval:</b> {interval['code']} ({interval['label']})
    </div>
    """, unsafe_allow_html=True)


# ═══════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════

st.markdown("# 🌾 PRF Live Rainfall Tracker")
st.markdown(
    f"**{interval['name']} 2026 · Interval {interval['code']}** "
    f"· Real-time CPC rainfall vs 3-year implied normals"
)
st.divider()


# ═══════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════

if generate or generate_report:

    date_start = f"2026-{interval['start']}"
    date_end   = f"2026-{interval['end']}"
    total_days = interval["total_days"]
    interval_name = interval["name"]

    rain_df = load_rainfall(date_start, date_end)
    tracker = build_tracker(grids_df, rain_df, coverage_level, total_days)

    # ─────────────────────────────────────
    # CLIENT MODE
    # ─────────────────────────────────────
    if client_mode and client_positions_df is not None and not client_positions_df.empty:

        # Get the unique grid IDs from client positions
        client_grid_ids = client_positions_df['GRID_ID'].unique().tolist()
        display_df = tracker[tracker["GRID_ID"].isin(client_grid_ids)].copy()

        if display_df.empty:
            st.warning("No rainfall data found for client grids. Check pipeline status.")
            st.stop()

        display_df = display_df.sort_values("PROJECTED_INDEX", ascending=True)
        days_in = int(display_df["DAYS_COLLECTED"].iloc[0]) if len(display_df) > 0 else 0
        likely_ct = len(display_df[display_df["SIGNAL"] == "LIKELY INDEMNITY"])

        # ── Metrics Row ──
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Positions", len(client_positions_df))
        c2.metric("Days", f"{days_in} / {total_days}")
        c3.metric("Coverage", f"{coverage_level}%")
        c4.metric("Likely Indemnity", likely_ct)

        # ──────────────────────────────────
        # INDEMNITY ESTIMATES
        # ──────────────────────────────────
        st.divider()
        st.markdown("### 💰 Portfolio Indemnity Estimates")

        indem_df = compute_indemnity_scenarios(
            client_positions_df, rain_df, grids_df, interval_name, total_days
        )

        if not indem_df.empty:
            total_pessimistic = int(indem_df['INDEM_PESSIMISTIC'].sum())
            total_trend       = int(indem_df['INDEM_TREND'].sum())
            total_optimistic  = int(indem_df['INDEM_OPTIMISTIC'].sum())

            # ── Scenario Cards ──
            sc1, sc2, sc3 = st.columns(3)
            with sc1:
                st.markdown(f"""
                <div class="indemnity-card" style="background:{FC_GREEN};">
                    <div style="font-size:13px; opacity:0.9;">OPTIMISTIC</div>
                    <div style="font-size:11px; opacity:0.7;">Zero rain remaining</div>
                    <div style="font-size:32px; margin-top:8px;">${total_optimistic:,.0f}</div>
                </div>
                """, unsafe_allow_html=True)
            with sc2:
                st.markdown(f"""
                <div class="indemnity-card" style="background:{FC_AMBER};">
                    <div style="font-size:13px; opacity:0.9;">CURRENT TREND</div>
                    <div style="font-size:11px; opacity:0.7;">Linear extrapolation</div>
                    <div style="font-size:32px; margin-top:8px;">${total_trend:,.0f}</div>
                </div>
                """, unsafe_allow_html=True)
            with sc3:
                st.markdown(f"""
                <div class="indemnity-card" style="background:{FC_RUST};">
                    <div style="font-size:13px; opacity:0.9;">PESSIMISTIC</div>
                    <div style="font-size:11px; opacity:0.7;">Normal pace remaining</div>
                    <div style="font-size:32px; margin-top:8px;">${total_pessimistic:,.0f}</div>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("")

            # ── Position Breakdown Table ──
            st.markdown("#### Position Breakdown")

            # Build pre-formatted string table to prevent Streamlit arrow indicators
            tbl = pd.DataFrame()
            tbl['Grid']       = indem_df['GRID_ID'].astype(str)
            tbl['County']     = indem_df['COUNTY']
            tbl['Acres']      = indem_df['ACRES'].apply(lambda x: f"{x:,}")
            tbl['Alloc %']    = indem_df['ALLOC_PCT'].apply(lambda x: f"{x:.0f}%")
            tbl['CBV']        = indem_df['CBV'].apply(lambda x: f"${x:.2f}")
            tbl['Normal']     = indem_df['NORMAL_IN'].apply(lambda x: f"{x:.2f}\"")
            tbl['Rain']       = indem_df['RAIN_SO_FAR'].apply(lambda x: f"{x:.2f}\"")
            tbl['Days']       = indem_df['DAYS'].astype(str)
            tbl['Cur Idx']    = indem_df['PARTIAL_INDEX'].apply(lambda x: f"{x:.1f}")
            tbl['Opt Idx']    = indem_df['OPTIMISTIC_IDX'].apply(lambda x: f"{x:.1f}")
            tbl['Trend Idx']  = indem_df['TREND_IDX'].apply(lambda x: f"{x:.1f}")
            tbl['Pess Idx']   = indem_df['PESSIMISTIC_IDX'].apply(lambda x: f"{x:.1f}")
            tbl['Protection'] = indem_df['PROTECTION'].apply(lambda x: f"${x:,.0f}")
            tbl['$ Optimistic']  = indem_df['INDEM_OPTIMISTIC'].apply(lambda x: f"${x:,.0f}")
            tbl['$ Trend']       = indem_df['INDEM_TREND'].apply(lambda x: f"${x:,.0f}")
            tbl['$ Pessimistic'] = indem_df['INDEM_PESSIMISTIC'].apply(lambda x: f"${x:,.0f}")

            # Totals row
            totals = pd.DataFrame([{
                'Grid': 'TOTAL',
                'County': '',
                'Acres': f"{int(indem_df['ACRES'].sum()):,}",
                'Alloc %': '',
                'CBV': '',
                'Normal': '',
                'Rain': '',
                'Days': '',
                'Cur Idx': '',
                'Opt Idx': '',
                'Trend Idx': '',
                'Pess Idx': '',
                'Protection': f"${int(indem_df['PROTECTION'].sum()):,}",
                '$ Optimistic': f"${int(indem_df['INDEM_OPTIMISTIC'].sum()):,}",
                '$ Trend': f"${int(indem_df['INDEM_TREND'].sum()):,}",
                '$ Pessimistic': f"${int(indem_df['INDEM_PESSIMISTIC'].sum()):,}",
            }])
            tbl = pd.concat([tbl, totals], ignore_index=True)

            st.dataframe(
                tbl, use_container_width=True, hide_index=True,
                height=min(600, 50 + len(tbl) * 40),
            )
        else:
            st.info(f"No client positions have allocation for {interval_name}.")

        # ── Gauges ──
        st.divider()
        st.markdown("### 📊 Projected Final Index")

        st.markdown(f"""
        <div style="
            display: flex; align-items: center; gap: 32px; 
            padding: 12px 20px; background: {FC_CREAM}; 
            border: 1px solid #d5d0c6; border-radius: 8px; margin-bottom: 16px;
            flex-wrap: wrap;
        ">
            <div style="display:flex; align-items:center; gap:8px;">
                <div style="width:40px; height:14px; background:{FC_GREEN}; border-radius:3px;"></div>
                <span style="font-size:13px; color:#2d3a2e;"><b>Bar</b> — Projected Final Index</span>
            </div>
            <div style="display:flex; align-items:center; gap:8px;">
                <div style="width:4px; height:22px; background:#2d3a2e; border-radius:1px;"></div>
                <span style="font-size:13px; color:#2d3a2e;"><b>Line</b> — Current Estimated Index Value</span>
            </div>
            <div style="display:flex; align-items:center; gap:8px;">
                <div style="width:40px; height:14px; background:rgba(94,151,50,0.28); border:1px solid #ccc; border-radius:3px;"></div>
                <span style="font-size:13px; color:#2d3a2e;">Indemnity zone (below {coverage_level})</span>
            </div>
        </div>
        """, unsafe_allow_html=True)

        for _, row in display_df.iterrows():
            fig = create_gauge(
                grid_id=row["GRID_ID"],
                projected_index=row["PROJECTED_INDEX"],
                partial_index=row["PARTIAL_INDEX"],
                signal=row["SIGNAL"],
                rain_so_far=row["RAIN_SO_FAR"],
                normal_in=row["NORMAL_IN"],
                days=row["DAYS_COLLECTED"],
                total_days=total_days,
                coverage_level=coverage_level,
                county_name=row.get("COUNTY_NAME"),
            )
            st.plotly_chart(fig, use_container_width=True)

            sig = row["SIGNAL"]
            proj = row["PROJECTED_INDEX"]
            part = row["PARTIAL_INDEX"]
            if sig == "LIKELY INDEMNITY":
                st.markdown(
                    f'<div class="signal-indemnity">'
                    f'✅ LIKELY INDEMNITY — Current: {part:.1f}  ·  Projected: {proj:.1f}  ·  Trigger: {coverage_level}'
                    f'</div>', unsafe_allow_html=True)
            else:
                st.markdown(
                    f'<div class="signal-ok">'
                    f'OK — Current: {part:.1f}  ·  Projected: {proj:.1f}'
                    f'</div>', unsafe_allow_html=True)

            st.markdown("")

        # ── Word Report Generation ──
        if generate_report and not indem_df.empty:
            with st.spinner("Generating Word report..."):
                all_positions = load_client_policy(client_name, crop_year)
                report_buf = build_word_report(
                    client_name=client_name,
                    crop_year=crop_year,
                    interval_name=interval_name,
                    interval_code=interval['code'],
                    interval_label=interval['label'],
                    coverage_level=coverage_level,
                    total_days=total_days,
                    indem_df=indem_df,
                    display_df=display_df,
                    all_positions_df=all_positions,
                    estimated_premium=647054,
                )
                filename = f"PRF_Report_{client_name.replace(' ', '_')}_{interval_name}_{crop_year}.docx"
                st.success("Report generated!")
                st.download_button(
                    label="⬇️ Download Report",
                    data=report_buf,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

    # ─────────────────────────────────────
    # MANUAL MODE (existing behavior)
    # ─────────────────────────────────────
    else:
        display_df = None
        if grid_entry and grid_entry.strip():
            try:
                typed_ids = [int(x.strip()) for x in grid_entry.split(",") if x.strip()]
                display_df = tracker[tracker["GRID_ID"].isin(typed_ids)].copy()
            except ValueError:
                st.error("Invalid grid IDs. Use comma-separated numbers like: 7929, 8230")
                st.stop()
        elif selected_labels:
            selected_ids = [label_to_id[lbl] for lbl in selected_labels]
            display_df = tracker[tracker["GRID_ID"].isin(selected_ids)].copy()
        elif show_all_likely:
            display_df = tracker[tracker["SIGNAL"] == "LIKELY INDEMNITY"].copy()
        else:
            n = len(tracker) if top_n == "All" else int(top_n)
            display_df = tracker.nsmallest(n, "PROJECTED_INDEX").copy()

        if display_df is None or display_df.empty:
            st.warning("No grids found.")
            st.stop()

        display_df = display_df.sort_values("PROJECTED_INDEX", ascending=True)

        days_in = int(display_df["DAYS_COLLECTED"].iloc[0])
        likely_ct = len(display_df[display_df["SIGNAL"] == "LIKELY INDEMNITY"])

        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Grids", len(display_df))
        c2.metric("Days", f"{days_in} / {total_days}")
        c3.metric("Coverage", f"{coverage_level}%")
        c4.metric("Likely Indemnity", likely_ct)

        st.divider()
        st.markdown("### 📊 Projected Final Index")

        # ── Legend ──
        st.markdown(f"""
        <div style="
            display: flex; align-items: center; gap: 32px; 
            padding: 12px 20px; background: {FC_CREAM}; 
            border: 1px solid #d5d0c6; border-radius: 8px; margin-bottom: 16px;
            flex-wrap: wrap;
        ">
            <div style="display:flex; align-items:center; gap:8px;">
                <div style="width:40px; height:14px; background:{FC_GREEN}; border-radius:3px;"></div>
                <span style="font-size:13px; color:#2d3a2e;"><b>Bar</b> — Projected Final Index</span>
            </div>
            <div style="display:flex; align-items:center; gap:8px;">
                <div style="width:4px; height:22px; background:#2d3a2e; border-radius:1px;"></div>
                <span style="font-size:13px; color:#2d3a2e;"><b>Line</b> — Current Estimated Index Value</span>
            </div>
            <div style="display:flex; align-items:center; gap:8px;">
                <div style="width:40px; height:14px; background:rgba(94,151,50,0.28); border:1px solid #ccc; border-radius:3px;"></div>
                <span style="font-size:13px; color:#2d3a2e;">Indemnity zone (below {coverage_level})</span>
            </div>
        </div>
        """, unsafe_allow_html=True)

        for _, row in display_df.iterrows():
            fig = create_gauge(
                grid_id=row["GRID_ID"],
                projected_index=row["PROJECTED_INDEX"],
                partial_index=row["PARTIAL_INDEX"],
                signal=row["SIGNAL"],
                rain_so_far=row["RAIN_SO_FAR"],
                normal_in=row["NORMAL_IN"],
                days=row["DAYS_COLLECTED"],
                total_days=total_days,
                coverage_level=coverage_level,
                county_name=row.get("COUNTY_NAME"),
            )
            st.plotly_chart(fig, use_container_width=True)

            sig = row["SIGNAL"]
            proj = row["PROJECTED_INDEX"]
            part = row["PARTIAL_INDEX"]
            if sig == "LIKELY INDEMNITY":
                st.markdown(
                    f'<div class="signal-indemnity">'
                    f'✅ LIKELY INDEMNITY — Current: {part:.1f}  ·  Projected: {proj:.1f}  ·  Trigger: {coverage_level}'
                    f'</div>', unsafe_allow_html=True)
            else:
                st.markdown(
                    f'<div class="signal-ok">'
                    f'OK — Current: {part:.1f}  ·  Projected: {proj:.1f}'
                    f'</div>', unsafe_allow_html=True)

            st.markdown("")

        st.divider()
        st.markdown("### 📋 Detail")

        table_df = display_df[[
            "GRID_ID", "COUNTY_NAME", "NORMAL_IN", "DAYS_COLLECTED",
            "RAIN_SO_FAR", "PARTIAL_INDEX", "PROJECTED_RAIN",
            "PROJECTED_INDEX", "SIGNAL", "CV_PCT"
        ]].copy()
        table_df.columns = [
            "Grid", "Counties", "Normal (in)", "Days",
            "Rain (in)", "Current Idx", "Proj Rain (in)",
            "Proj Index", "Signal", "CV%"
        ]

        st.dataframe(
            table_df, use_container_width=True, hide_index=True,
            height=min(600, 50 + len(table_df) * 40),
            column_config={
                "Grid": st.column_config.NumberColumn(format="%d"),
                "Normal (in)": st.column_config.NumberColumn(format="%.1f"),
                "Rain (in)": st.column_config.NumberColumn(format="%.1f"),
                "Proj Rain (in)": st.column_config.NumberColumn(format="%.1f"),
                "Current Idx": st.column_config.NumberColumn(format="%.1f"),
                "Proj Index": st.column_config.NumberColumn(format="%.1f"),
                "CV%": st.column_config.NumberColumn(format="%.1f"),
            }
        )

else:
    st.markdown(f"""
    <div style='text-align:center; padding:80px 20px; color:{FC_SLATE};'>
        <h2 style='color:{FC_GREEN};'>Select grids and click Generate</h2>
        <p style='font-size:16px;'>
            Choose your interval and coverage level.<br>
            Load a client policy or select grids manually.<br>
            Then click Generate.
        </p>
        <p style='font-size:14px; margin-top:20px;'>
            🌾 Insured grids · 3-year implied normals · 0.4% avg CV · RT data through today
        </p>
    </div>
    """, unsafe_allow_html=True)
